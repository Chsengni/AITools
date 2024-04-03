import tkinter as tk
from tkinter import ttk
from tkinter import filedialog, messagebox
import threading
from docx import Document
import erniebot
import os
import queue
from openai import OpenAI
from dashscope import Generation
from http import HTTPStatus
from docx.shared import Pt
import json
import dashscope

# 设置默认值
default_access_token = ''
default_prompt = [
    "你是一个文本降重机器，你只执行降重改写 语序颠倒 顺序调换 同义替换 句子意思不变 主动句改被动句 被动句改主动句 直接输出结果",
    "你是一个文本润色机器，你只执行润色文本使其更流畅、更具吸引力同时保留原意 句子意思不变 直接输出结果",
    "你是一个文本校对机器，你只执行校对文字 正确表达句子 不能出现逻辑错误 句子意思不变 直接输出结果"
]
default_min_paragraph_length = '100'

def call_chatgpt_thread(api_key, model, prompt, text, result_queue):
    try:
        client = OpenAI(
            api_key=api_key,
            base_url="https://chatgpt.24z.cn/v1"
        )
        completion = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": text}
            ]
        )
        response = completion.choices[0].message.content
        result_queue.put(response.replace(" ", "").replace("\n", ""))
    except Exception as e:
        result_queue.put(e)

def call_erniebot_thread(api_key, model, prompt, text, result_queue):
    erniebot.api_type = 'aistudio'
    erniebot.access_token = api_key
    try:
        response_stream = erniebot.ChatCompletion.create(
            model=model,
            messages=[{
                'role': 'user',
                'content': text
            }],
            system=prompt,
            stream=True
        )
        result_queue.put(''.join([response.get_result() for response in response_stream]).replace(" ", "").replace("\n", ""))
    except Exception as e:
        result_queue.put(e)

def call_qwen_thread(api_key, model, prompt, text, result_queue):
    try:
        dashscope.api_key =api_key
        if model not in ["qwen-1.8b-chat", "qwen-72b-chat", "qwen1.5-72b-chat", "qwen1.5-14b-chat", "qwen1.5-7b-chat",
                         "qwen-14b-chat", "qwen-7b-chat", "qwen-1.8b-longcontext-chat"]:
            gen = Generation.call(
                model=model,
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": text}
                ],
                result_format='message',
                enable_search=True,
                stream=True,
                incremental_output=True
            )
        else:
            gen = Generation.call(
                model=model,
                messages=[
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": text}
                ],
                result_format='message',
                stream=True,
                incremental_output=True
            )
        resp = ""
        for response in gen:
            if response.status_code == HTTPStatus.OK:
                response = response.output.choices[0].message.content
                resp += response
        response = resp

        result_queue.put(response.replace(" ", "").replace("\n", ""))
    except Exception as e:
        result_queue.put(e)

def call_ai(selected_option, api_key, model, prompt, text):
    result_queue = queue.Queue()
    if selected_option == "文心一言":
        thread = threading.Thread(target=call_erniebot_thread, args=(api_key, model, prompt, text, result_queue))
        thread.start()
        return str(result_queue.get())
    elif selected_option == "通义千问":
        thread = threading.Thread(target=call_qwen_thread, args=(api_key, model, prompt, text, result_queue))
        thread.start()
        return str(result_queue.get())
    else:
        thread = threading.Thread(target=call_chatgpt_thread, args=(api_key, model, prompt, text, result_queue))
        thread.start()
        return str(result_queue.get())

def process_file():
    selected_option = dropdown.get()
    access_token_val = access_token.get().strip()
    model_val = model_dropdown.get().strip()
    prompt_val = prompt_var.get("1.0", tk.END).strip()
    min_paragraph_length_val = min_paragraph_length_var.get().strip()
    file_path_val = input_file_path_var.get().strip()
    
    try:
        file_path_val = eval(file_path_val)
    except:
        pass

    if isinstance(file_path_val, str):
        output_path = output_file_path_var.get().strip()
        if not all([access_token_val, model_val, prompt_val, min_paragraph_length_val, file_path_val, output_path]):
            messagebox.showwarning("提示", "不能为空")
        else:
            min_paragraph_length = int(min_paragraph_length_val)
            doc = Document(file_path_val)

            if not any(style.name == 'Normal' for style in doc.styles):
                doc_styles = doc.styles
                new_style = doc_styles.add_style('Normal', 1)
                new_style.font.name = '宋体'
                new_style.font.size = Pt(12)
            paragraphs = doc.paragraphs[:]
            tables = doc.tables[:]
            total_elements = len(paragraphs) + len(tables)
            progress['maximum'] = total_elements

            for i, para in enumerate(paragraphs, start=1):
                text = para.text.strip()
                if len(text) > min_paragraph_length:
                    api_result = call_ai(selected_option, access_token_val, model_val, prompt_val, text)
                    if api_result:
                        para.clear()
                        para.add_run(api_result)
                    else:
                        messagebox.showerror("提示", "网络异常")
                progress['value'] = i
                root.update_idletasks()

            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            text = para.text.strip()
                            if len(text) > min_paragraph_length:
                                api_result = call_ai(selected_option, access_token_val, model_val, prompt_val, text)
                                if api_result:
                                    para.clear()
                                    para.add_run(api_result)
                                else:
                                    messagebox.showerror("提示", "网络异常")
                            progress['value'] = i
                            root.update_idletasks()
                            i += 1

            output_file_path = os.path.join(output_path, file_path_val.split("/")[-1][:-5] + "_" + ability_dropdown.get() + ".docx")
            doc.save(output_file_path)
            progress['value'] = 0
            root.update_idletasks()
            messagebox.showinfo("完成", f"已保存为 {output_file_path}")

    if isinstance(file_path_val, tuple):
        for file in file_path_val:
            output_path = output_file_path_var.get().strip()
            if not all([access_token_val, model_val, prompt_val, min_paragraph_length_val, file_path_val, output_path]):
                messagebox.showwarning("提示", "不能为空")
            else:
                min_paragraph_length = int(min_paragraph_length_val)
                doc = Document(file)

                if not any(style.name == 'Normal' for style in doc.styles):
                    doc_styles = doc.styles
                    new_style = doc_styles.add_style('Normal', 1)
                    new_style.font.name = '宋体'
                    new_style.font.size = Pt(12)

                paragraphs = doc.paragraphs[:]
                tables = doc.tables[:]
                total_elements = len(paragraphs) + len(tables)
                progress['maximum'] = total_elements

                for i, para in enumerate(paragraphs, start=1):
                    text = para.text.strip()
                    if len(text) > min_paragraph_length:
                        api_result = call_ai(selected_option, access_token_val, model_val, prompt_val, text)
                        if api_result:
                            para.clear()
                            para.add_run(api_result)
                        else:
                            messagebox.showerror("提示", "网络异常")
                    progress['value'] = i
                    root.update_idletasks()

                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                text = para.text.strip()
                                if len(text) > min_paragraph_length:
                                    api_result = call_ai(selected_option, access_token_val, model_val, prompt_val, text)
                                    if api_result:
                                        para.clear()
                                        para.add_run(api_result)
                                    else:
                                        messagebox.showerror("提示", "网络异常")
                                progress['value'] = i
                                root.update_idletasks()
                                i += 1

                output_file_path = os.path.join(output_path, file.split("/")[-1][:-5] + "_" + ability_dropdown.get() + ".docx")
                doc.save(output_file_path)
                progress['value'] = 0
                root.update_idletasks()
                messagebox.showinfo("完成", f"已保存为 {output_file_path}")
                
                
def execute_thread():
    thread = threading.Thread(target=process_file)
    thread.start()


def choose_input_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word文档", "*.docx")], title="选择输入文件")
    if file_path:
        input_file_path_var.set(file_path)


def select_output_folder():
    folder_path = filedialog.askdirectory(title="选择输出文件夹")
    if folder_path:
        output_file_path_var.set(folder_path)


def change_ability(event):
    if ability_dropdown.get() == "降重":
        prompt_var.delete('1.0', tk.END)
        prompt_var.insert(tk.END, default_prompt[0])
        execute_button['text'] = "开始降重"
    elif ability_dropdown.get() == "润色":
        prompt_var.delete('1.0', tk.END)
        prompt_var.insert(tk.END, default_prompt[1])
        execute_button['text'] = "开始润色"
    elif ability_dropdown.get() == "校对":
        prompt_var.delete('1.0', tk.END)
        prompt_var.insert(tk.END, default_prompt[2])
        execute_button['text'] = "开始校对"
    else:
        selected_function = ability_dropdown.get()
        if selected_function in config:
            prompt_var.delete('1.0', tk.END)
            prompt_var.insert(tk.END, config[selected_function])
        else:
            prompt_var.delete('1.0', tk.END)
            prompt_var.insert(tk.END, "未定义角色")

def update_ability_dropdown(ability_dropdown):
    # 更新功能下拉框中的选项
    ability_choices = ['降重', '润色', '校对']  # 默认选项
    custom_functions = list(config.keys())
    if custom_functions:
        ability_choices.extend(custom_functions)
    ability_dropdown['values'] = ability_choices

def change_model(event):
    # 创建左侧下拉框
    selected_option = dropdown.get()
    if selected_option == "文心一言":
        model_dropdown["value"] = ["ernie-3.5", "ernie-turbo", "ernie-4", "ernie-longtext"]
        model_dropdown.current(0)
        
    elif selected_option == "通义千问":
        model_dropdown["value"] = ["qwen-1.8b-chat", "qwen-72b-chat", "qwen1.5-72b-chat", "qwen1.5-14b-chat",
                                    "qwen1.5-7b-chat", "qwen-14b-chat", "qwen-7b-chat", "qwen-1.8b-longcontext-chat",
                                    "qwen-turbo", "qwen-plus", "qwen-max", "qwen-max-1201", "qwen-max-longcontext"]
        model_dropdown.current(0)

    else:
        model_dropdown["value"] = ["gpt-3.5-turbo"]
        model_dropdown.current(0)
    api_key = load_apikey()
    if api_key!={}:
        access_token.delete(0, tk.END)
        access_token.insert(0, api_key[selected_option])
def open_folder():
    folder_path = filedialog.askopenfilenames(title="打开文件夹")
    if folder_path:
        input_file_path_var.set(folder_path)

def about():
    messagebox.showinfo("关于", "作者：Chsengni\n联系方式QQ：2371778707\n版本：1.0.0\nhttps://aistudio.baidu.com/serving/app/13662/")

def save_config(config):
    # 保存配置到文件
    with open('config.json', 'w') as f:
        json.dump(config, f)

def load_config():
    # 加载配置文件
    try:
        with open('config.json', 'r') as f:
            config = json.load(f)
    except FileNotFoundError:
        config = {}
    return config


def add_custom_function():
    custom_function_window = tk.Toplevel(root)
    custom_function_window.iconbitmap('favicon.ico')   # 更改窗口图标
    # 设置窗口居中
    window_width = 250
    window_height = 150
    screen_width = custom_function_window.winfo_screenwidth()
    screen_height = custom_function_window.winfo_screenheight()
    x_coordinate = int((screen_width / 2) - (window_width / 2))
    y_coordinate = int((screen_height / 2) - (window_height / 2))
    custom_function_window.geometry(f'{window_width}x{window_height}+{x_coordinate}+{y_coordinate}')
    custom_function_window.resizable(False, False)  # Disable resizing

    custom_function_window.title("自定义功能")
    custom_functions = list(config.keys())
    # 添加功能名称输入框
    functions_label = tk.Label(custom_function_window, text="已有功能")
    functions_label.grid(row=0, column=0, padx=5, pady=5)
    role_combobox = ttk.Combobox(custom_function_window, values=custom_functions)
    role_combobox.grid(row=0, column=1,columnspan=3, padx=5, pady=5)
    if custom_functions!=[]:
        role_combobox.current(0)
    # 添加功能名称输入框
    function_name_label = tk.Label(custom_function_window, text="功能名称")
    function_name_label.grid(row=1, column=0, padx=5, pady=5)
    
    if custom_functions!=[] and role_combobox.get() != "":
        function_name_entry = tk.Entry(custom_function_window,textvariable=role_combobox.get())
    else:
        function_name_entry = tk.Entry(custom_function_window)
    function_name_entry.grid(row=1, column=1,columnspan=3, padx=5, pady=5)

    # 添加角色输入框
    role_label = tk.Label(custom_function_window, text="角色")
    role_label.grid(row=2, column=0, padx=5, pady=5)
    if custom_functions!=[] and role_combobox.get() != "":
        role_entry = tk.Entry(custom_function_window,textvariable=config[role_combobox.get()])
    else:
        role_entry = tk.Entry(custom_function_window)
    role_entry.grid(row=2, column=1,columnspan=3,padx=5, pady=5)
    # 加载角色按钮
    def load_role(event=None):
        selected_function = role_combobox.get()
        if selected_function in config:
            function_name_entry.delete(0, tk.END)
            function_name_entry.insert(0, selected_function)
            role_entry.delete(0, tk.END)
            role_entry.insert(0, config[selected_function])
    load_role()
    role_combobox.bind("<<ComboboxSelected>>", load_role)

    # 保存按钮
    def save_function():
        function_name = function_name_entry.get()
        role = role_entry.get()
        if function_name and role:
            config[function_name] = role
            save_config(config)
            messagebox.showinfo("保存成功", "功能已保存")
            update_ability_dropdown(ability_dropdown)  # 初始化下拉框选项
            custom_function_window.destroy()
        else:
            messagebox.showerror("保存失败", "功能和角色不能为空")
            add_custom_function()

    save_button = tk.Button(custom_function_window, text="保存配置", command=save_function)
    save_button.grid(row=3, column=1,columnspan=1, padx=5, pady=5)

    # 移除按钮
    def remove_function():
        function_name = function_name_entry.get()
        if function_name:
            if function_name in config:
                del config[function_name]
                save_config(config)
                messagebox.showinfo("移除成功", "功能已移除")
                update_ability_dropdown(ability_dropdown)  # 初始化下拉框选项
                custom_function_window.destroy()
            else:
                messagebox.showerror("移除失败", "该功能不存在")
                add_custom_function()
        else:
            messagebox.showerror("移除失败", "请输入功能")
            add_custom_function()

    remove_button = tk.Button(custom_function_window, text="移除配置", command=remove_function)
    remove_button.grid(row=3, column=3, padx=5, pady=5)

def save_apikeys(api_keys):
    with open("api_keys.json", "w") as file:
        json.dump(api_keys, file)


def load_apikey():
    # 加载配置文件
    try:
        with open('api_keys.json', 'r') as f:
            api_keys = json.load(f)
    except FileNotFoundError:
        api_keys = {}
    return api_keys


def add_apikey():
    # 创建一个新的顶级窗口
    apikey_window = tk.Toplevel(root)
    apikey_window.title("添加令牌")
    apikey_window.iconbitmap('favicon.ico')   # 更改窗口图标
    window_width = 220
    window_height = 150
    screen_width = apikey_window.winfo_screenwidth()
    screen_height = apikey_window.winfo_screenheight()
    x_coordinate = int((screen_width / 2) - (window_width / 2))
    y_coordinate = int((screen_height / 2) - (window_height / 2))
    apikey_window.geometry(f'{window_width}x{window_height}+{x_coordinate}+{y_coordinate}')
    apikey_window.resizable(False, False)  # Disable resizing
    api_key = load_apikey()
    # 文心一言输入框和标签
    erniebot_label = tk.Label(apikey_window, text="文心一言")
    erniebot_label.grid(row=0, column=0, sticky='we', padx=5, pady=5)
    erniebot_entry = tk.Entry(apikey_window)
    erniebot_entry.grid(row=0, column=1, sticky='we', padx=5, pady=5, columnspan=1)

    # 通义千问输入框和标签
    qwen_label = tk.Label(apikey_window, text="通义千问")
    qwen_label.grid(row=1, column=0, sticky='we', padx=5, pady=5)
    qwen_entry = tk.Entry(apikey_window)
    qwen_entry.grid(row=1, column=1, sticky='we', padx=5, pady=5, columnspan=1)

    # ChatGPT输入框和标签
    chatgpt_label = tk.Label(apikey_window, text="ChatGPT")
    chatgpt_label.grid(row=2, column=0, sticky='we', padx=5, pady=5)
    chatgpt_entry = tk.Entry(apikey_window)
    chatgpt_entry.grid(row=2, column=1, sticky='we', padx=5, pady=5, columnspan=1)

    def save_apikey():
        # 获取文心一言、通义千问和 ChatGPT 的 API 密钥输入框中的值
        erniebot_key = erniebot_entry.get()
        qwen_key = qwen_entry.get()
        chatgpt_key = chatgpt_entry.get()

        # 将 API 密钥信息保存到 JSON 文件中
        api_keys = {
            "文心一言": erniebot_key,
            "通义千问": qwen_key,
            "ChatGPT": chatgpt_key
        }

        for key in api_keys.keys():
            if key in api_key:
                del api_key[key]
        save_apikeys(api_keys)
        messagebox.showinfo("保存成功", "令牌已保存")
        
    if api_key!={}:
        chatgpt_entry.delete(0, tk.END)
        chatgpt_entry.insert(0, api_key['ChatGPT'])
        erniebot_entry.delete(0, tk.END)
        erniebot_entry.insert(0, api_key['文心一言'])
        qwen_entry.delete(0, tk.END)
        qwen_entry.insert(0, api_key['通义千问'])
        
    # 添加保存按钮
    save_button = tk.Button(apikey_window, text="保存", command=save_apikey)
    save_button.grid(row=3, column=0, columnspan=2, sticky='we', padx=5, pady=5)

root = tk.Tk()
root.title("AI文档处理工具")
root.iconbitmap('favicon.ico')   # 更改窗口图标
# 设置窗口居中
window_width = 550
window_height = 350
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
x_coordinate = int((screen_width / 2) - (window_width / 2))
y_coordinate = int((screen_height / 2) - (window_height / 2))
root.geometry(f'{window_width}x{window_height}+{x_coordinate}+{y_coordinate}')

root.resizable(False, False)  # Disable resizing

# 创建菜单栏
menubar = tk.Menu(root)

# 文件菜单
filemenu = tk.Menu(menubar, tearoff=0)
filemenu.add_command(label="打开文件", command=choose_input_file)
filemenu.add_command(label="打开文件夹", command=open_folder)

filemenu.add_command(label="选择输出文件夹", command=select_output_folder)
menubar.add_cascade(label="文件", menu=filemenu)

# 配置菜单
configmenu = tk.Menu(menubar, tearoff=0)
configmenu.add_command(label="自定义功能", command=add_custom_function)
configmenu.add_command(label="令牌配置", command=add_apikey)
menubar.add_cascade(label="配置", menu=configmenu)
# 加载配置文件
config = load_config()
# 关于菜单
aboutmenu = tk.Menu(menubar, tearoff=0)
aboutmenu.add_command(label="关于", command=about)
menubar.add_cascade(label="关于", menu=aboutmenu)

# 将菜单栏应用到窗口
root.config(menu=menubar)

# 表格布局
root.grid_rowconfigure(0, weight=1)
root.grid_columnconfigure(0, weight=1)
root.grid_columnconfigure(1, weight=1)
root.grid_columnconfigure(2, weight=1)
root.grid_columnconfigure(3, weight=1)

# 平台标签和下拉框
machine_label = tk.Label(root, text="平台")
machine_label.grid(row=0, column=0, padx=10)
options = ["文心一言", "通义千问", "ChatGPT"]
dropdown = ttk.Combobox(root, values=options)
dropdown.bind("<<ComboboxSelected>>", change_model)
dropdown.current(0)
dropdown.grid(row=0, column=1, sticky='we', padx=5, pady=5, columnspan=1)

# 模型标签和下拉框
model_label = tk.Label(root, text="模型")
model_label.grid(row=0, column=2, sticky='we')
model_options = ["ernie-3.5", "ernie-turbo", "ernie-4", "ernie-longtext"]
model_dropdown = ttk.Combobox(root, values=model_options)
model_dropdown.current(0)
model_dropdown.grid(row=0, column=3, sticky='we', padx=5, pady=5, columnspan=1)

# 令牌标签和输入框
access_token_label = tk.Label(root, text="令牌")
access_token_label.grid(row=1, column=0, sticky='we')
access_token = tk.Entry(root)
api_key= load_apikey()
if api_key!={}:
    access_token.insert(tk.END,api_key[dropdown.get()] )
else:
    access_token.insert(tk.END, default_access_token)
access_token.grid(row=1, column=1, sticky='we', padx=5, pady=5, columnspan=1)


# 功能标签和下拉框
ability_label = tk.Label(root, text="功能")
ability_label.grid(row=1, column=2, sticky='we')
#ability_choices = ['降重', '润色', '校对']
ability_dropdown = ttk.Combobox(root)
update_ability_dropdown(ability_dropdown)  # 初始化下拉框选项
ability_dropdown.bind("<<ComboboxSelected>>", change_ability)
ability_dropdown.current(0)
ability_dropdown.grid(row=1, column=3, sticky='we', padx=5, pady=5, columnspan=1)

# 提示标签和文本框
prompt_label = tk.Label(root, text="角色")
prompt_label.grid(row=2, column=0, sticky='we')
prompt_var = tk.Text(root, height=6, wrap=tk.WORD)
prompt_var.insert(tk.END, default_prompt[0])
prompt_var.grid(row=2, column=1, columnspan=3, sticky='we', padx=5, pady=5)

# 最小段落长度标签和输入框
min_paragraph_length_label = tk.Label(root, text="最小段落长度")
min_paragraph_length_label.grid(row=3, column=0, sticky='we')
min_paragraph_length_var = tk.StringVar(value=default_min_paragraph_length)
min_paragraph_length_entry = tk.Spinbox(root, from_=0, to=1000, textvariable=min_paragraph_length_var)
min_paragraph_length_entry.grid(row=3, column=1, sticky='we', padx=5, pady=5, columnspan=3)

# 选择输入文件按钮和输入框
input_file_button = tk.Button(root, text="选择输入文件", command=choose_input_file)
input_file_button.grid(row=4, column=0, sticky='we', padx=5, pady=5)
input_file_path_var = tk.StringVar()
input_file_path_entry = tk.Entry(root, textvariable=input_file_path_var, state='readonly')
input_file_path_entry.grid(row=4, column=1, sticky='we', padx=5, pady=5, columnspan=1)

# 选择输出文件夹按钮和输入框
output_file_button = tk.Button(root, text="选择输出文件夹", command=select_output_folder)
output_file_button.grid(row=4, column=2, sticky='we', padx=5, pady=5)
output_file_path_var = tk.StringVar()
output_file_path_entry = tk.Entry(root, textvariable=output_file_path_var, state='readonly')
output_file_path_entry.grid(row=4, column=3, sticky='we', padx=5, pady=5, columnspan=1)

# 开始执行按钮
execute_button = tk.Button(root, text="开始" + ability_dropdown.get(), command=execute_thread)
execute_button.grid(row=6, column=0, columnspan=4, sticky='we', padx=5, pady=5)

# 进度条
progress = ttk.Progressbar(root, orient="horizontal", length=1000)
progress.grid(row=7, column=0, columnspan=4, pady=20, padx=10, sticky='we')

root.mainloop()
